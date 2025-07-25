from typing import List, Optional
from llama_index.core.schema import Document
import os

try:
    from docx import Document as DocxDocument
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

def load_docx_as_documents(file_path: Optional[str] = None, text: Optional[str] = None, file_obj=None) -> List[Document]:
    """
    Load a .docx file from the local filesystem, a file-like object, or accept a string, returning a list of LlamaIndex Document objects.

    Args:
        file_path: Path to the .docx file.
        text: Optional string to use as document content instead of loading from file.
        file_obj: Optional file-like object (e.g., from FastAPI UploadFile) to load the docx from memory.

    Returns:
        List of LlamaIndex Document objects.
    """
    if text is not None:
        return [Document(text=text)]
    if file_obj is not None:
        docx = DocxDocument(file_obj)
        full_text = []
        for para in docx.paragraphs:
            full_text.append(para.text)
        content = "\n".join(full_text)
        return [Document(text=content)]
    if file_path is None:
        raise ValueError("Either file_path, file_obj, or text must be provided.")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    docx = DocxDocument(file_path)
    full_text = []
    for para in docx.paragraphs:
        full_text.append(para.text)
    content = "\n".join(full_text)
    return [Document(text=content)]  